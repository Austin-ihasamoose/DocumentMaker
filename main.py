'''
0: Line 1: Shop Log for TEMPLATE to TEMPLATE
    - Style: ARIAL 14, BOLD, UNDERLINE, ITALIC
2: DATE (PERSON)
    - Style: ARIAL 12, UNDERLINE
3: PERSON:
    - Style: ARIAL 12
8: DATE (PERSON)
9: PERSON:
14: DATE (PERSON)
15: PERSON:
20: DATE (PERSON)
21: PERSON:
26: DATE (PERSON)
27: PERSON:
'''

from pprint import pprint
from isoweek import Week
import os
import datetime
import calendar
import docx
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, Inches

ERRORS = []


#                 START DOCUMENT STYLE RELATED FUNCTIONS                  #


def add_styled_par(document, text, bold=False, underline=False, italic=False, font_name='Arial', font_size=12, spacing = WD_LINE_SPACING.SINGLE):
    # Add style options
    par = document.add_paragraph()  # Add new paragraph
    run = par.add_run(text)  # Run style settings
    if bold is True:
        run.bold = True
    if underline is True:
        run.underline = True
    if italic is True:
        run.italic = True
    run.font.name = font_name
    run.font.size = Pt(font_size)
    par.paragraph_format.line_spacing_rule = spacing


def add_bullet_par(document, text='', amount=1):
    for bullet in range(amount):
        par = document.add_paragraph()  # Add new paragraph
        run = par.add_run(text)  # Run style settings
        run.font.name = 'Arial'
        par.style = document.styles['List Bullet']
        par.paragraph_format.left_indent = Inches(0.5)
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


#                 END DOCUMENT STYLE RELATED FUNCTIONS                    #

#                 START DIRECTORY RELATED FUNCTIONS                       #


def shorthand_months():
    """
    Formats months and returns a dict in format
        MONTH: INT
        Ex.
        Jan: 1
    """

    sh_list = []
    for i in range(1, 13):
        month = calendar.month_name[i][0:3].capitalize()
        sh_list.append(month)
    return sh_list


def generate_folders(year):
    '''
    Generates folders for each year in "output/"
    Folder structure wanted:
    YEAR \\ MONTH \\
    :param year: Root year wanted
    '''

    printed = False
    year_path = 'output\\' + str(year)
    try:
        os.mkdir(year_path)
    except OSError as e:
        print(str(e))
    else:
        print("Successfully created the directory '%s'" % year_path)

    sh_list = shorthand_months()

    for i in range(1,13):
        month_path = 'output\\' + str(year) + '\\' + str(i) + ' - ' + str(sh_list[i-1])
        try:
            os.mkdir(month_path)
        except OSError as e:
            if printed is False:
                print("[ERROR]: MONTH FOLDERS ALREADY EXIST\n" + str(e))
            printed = True
        else:
            print("Successfully created the directory '%s'" % month_path)


#                 END DIRECTORY RELATED FUNCTIONS                        #

#                 START CALENDAR RELATED FUNCTIONS                       #

def generate_dates(year):
    # Structure: { month: { week_num: [from, to] } }

    c = calendar.Calendar()
    document_dates = {}

    for year_months in c.yeardatescalendar(year, width=1):
        for months in year_months:
            for week in months:
                month = week[0].month
                week_num = week[0].isocalendar()[1]
                start_day = week[0].day
                end_day = week[4].day
                document_dates[week_num] = [month, [start_day, end_day]]

    return document_dates


#                 END CALENDAR RELATED FUNCTIONS                       #

#                 START GENERATION RELATED FUNCTIONS                   #

def format_date(date):
    plain = str(date)
    year = plain[0:4]
    month = plain[5:7]
    day = plain[8:10]
    sh_list = shorthand_months()
    sh_reference = []

    for i in range(1, 13):
        sh_reference.append(i)

    for val in sh_reference:  # 1-12
        if int(month) == val:
            month = sh_list[val - 1]
            break

    return "{} {} {}".format(month, day, year)


def generate_document(year, iso_week, output_path):
    # week_start is iso week
    monday = datetime.date.fromisocalendar(year, iso_week, 1)
    friday = datetime.date.fromisocalendar(year, iso_week, 5)
    output_name = format_date(monday)[0:6] + ' to ' + format_date(friday)[0:6] + '.docx'

    if os.path.isfile(output_path+output_name):
        return print("[ERROR]: Already created: " + output_name + ", aborting")

    doc = docx.Document()
    sh_list = shorthand_months()

    add_styled_par(doc, "Shop Log for " + format_date(monday) + " to " + format_date(friday), bold=True, underline=True,
                   italic=True, font_size=14, spacing=WD_LINE_SPACING.ONE_POINT_FIVE)

    for i in range(1, 6):
        add_styled_par(doc,
                       format_date(datetime.date.fromisocalendar(year, iso_week, i))[0:6] + " (PERSON):",
                       underline=True, font_size=12)
        add_bullet_par(doc, amount=3)
        doc.add_paragraph()

    doc.save(output_path+output_name)

    return print('Generated: ' + output_name)


def generate_documents(year):
    sh_list = shorthand_months()
    results = generate_dates(year)
    c = calendar.Calendar()
    document_dates = {}
    first = True
    path = 'output\\' + str(year) + '\\'

    for i in range(len(sh_list)):
        if first is True:
            path += '1 - Jan\\'
            generate_document(year, 1, path)
            first = False

        for key in results.keys():
            if results[key][0] == i+1:
                path = 'output\\' + str(year) + '\\' + str(i+1) + ' - '+ sh_list[i] + '\\'
                generate_document(year, key, path)

    for year_months in c.yeardatescalendar(year, width=1):
        for months in year_months:
            for week in months:
                month = week[0].month
                week_num = week[0].isocalendar()[1]
                start_day = week[0].day
                end_day = week[4].day
                document_dates[week_num] = [month, [start_day, end_day]]

    return document_dates


#                 END GENERATION RELATED FUNCTIONS                   #


if __name__ == '__main__':
    year_to_generate = 2020
    generate_folders(year_to_generate)
    generate_documents(2020)